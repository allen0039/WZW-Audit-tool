#!/usr/bin/env python3
"""
全量提取：扫描 tiku 文件夹下所有文件（xlsx/docx），
提取全部数据为统一格式，不做任何过滤。
"""
import os, re, json, sys
from collections import OrderedDict
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR if (SCRIPT_DIR / 'tiku').exists() else SCRIPT_DIR.parent
BASE = PROJECT_ROOT / 'tiku' / '北京优度'
OUTPUT = PROJECT_ROOT / 'tiku_complete.json'
OUTPUT_MIN = PROJECT_ROOT / 'tiku_complete_min.json'

INDEX_PAT = re.compile(r'^[A-Z]{1,4}\d*(?:-\d+)*$')

def get_files():
    """Return list of (filepath, ext) for all files in BASE."""
    results = []
    for root, dirs, fnames in os.walk(BASE):
        for fn in sorted(fnames):
            if fn.startswith('~$'):
                continue
            fp = os.path.join(root, fn)
            ext = fn.rsplit('.', 1)[-1].lower()
            results.append((fp, ext))
    return results

def extract_index_from_text(text):
    """Try to find an index number in arbitrary text."""
    m = INDEX_PAT.search(text)
    return m.group(0) if m else None

def process_xlsx(fp, rel):
    """Extract ALL data from an xlsx file.
    Returns list of entry dicts."""
    import openpyxl
    entries = []
    try:
        wb = openpyxl.load_workbook(fp, data_only=True)
    except Exception as e:
        return [{'type': 'xlsx', 'source': rel, 'error': str(e)}]
    
    for sname in wb.sheetnames:
        ws = wb[sname]
        max_r = min(ws.max_row or 1000, 1000)
        max_c = min(ws.max_column or 30, 30)
        
        if max_r < 2:
            continue
        
        # Read ALL cells (0-indexed)
        all_cells = []
        for rr in ws.iter_rows(min_row=1, max_row=max_r, max_col=max_c, values_only=False):
            row = []
            for cell in rr:
                v = cell.value
                if v is not None:
                    s = str(v).strip()
                    if s:
                        row.append({'r': cell.row-1, 'c': cell.column-1, 'v': s})
            all_cells.append(row)
        
        # Strategy: split into tables by blank-row boundaries
        # A table starts after 2+ consecutive blank rows (or at row 0)
        tables = []
        table_start = 0
        blank_count = 0
        for ri, row in enumerate(all_cells):
            if not row:
                blank_count += 1
                if blank_count >= 2 and table_start < ri:
                    # End current table
                    tables.append((table_start, ri - blank_count + 1))
                    table_start = ri + 1
                    blank_count = 0
            else:
                blank_count = 0
        # Last table
        if table_start < len(all_cells):
            tables.append((table_start, len(all_cells)))
        
        for t_start, t_end in tables:
            if t_end - t_start < 2:
                continue
            # Collect cells for this table
            cells = []
            for ri in range(t_start, t_end):
                for cell in all_cells[ri]:
                    cells.append({'r': cell['r'] - t_start, 'c': cell['c'], 'v': cell['v']})
            
            if not cells:
                continue
            
            # Find index number: search first 6 rows, columns 5-15
            # Also search for "索引号：XXX" patterns anywhere in table
            idx = None
            title = ''
            for cell in cells:
                if cell['r'] <= 5 and 5 <= cell['c'] <= 15:
                    if INDEX_PAT.match(cell['v']):
                        idx = cell['v']
                        break
            # Fallback: search for cells with "索引号" near an index pattern
            if not idx:
                for cell in cells:
                    if '索引号' in cell['v']:
                        for c2 in cells:
                            if c2['r'] == cell['r'] and c2['c'] == cell['c'] + 1 and INDEX_PAT.match(c2['v']):
                                idx = c2['v']
                                break
                    if idx:
                        break
            
            # Also check sheet name for index
            if not idx:
                m2 = INDEX_PAT.match(sname.strip())
                if m2:
                    idx = m2.group(0)
            
            title = next((c['v'] for c in cells if c['c'] == 1 and c['r'] == 0), sname)
            
            entry = {
                'type': 'xlsx',
                'source': rel,
                'sheet': sname,
                'table_start_row': t_start,
                'table_end_row': t_end,
                'rows': t_end - t_start,
                'cols': max(c['c'] for c in cells) + 1 if cells else 0,
                'title': title,
                'cells': cells,
            }
            if idx:
                entry['index_no'] = idx
            
            entries.append(entry)
    
    wb.close()
    return entries

def process_docx(fp, rel):
    """Extract ALL data from a docx file.
    Returns list of entry dicts."""
    from docx import Document
    entries = []
    
    try:
        doc = Document(fp)
    except Exception as e:
        return [{'type': 'docx', 'source': rel, 'error': str(e)}]
    
    # Extract paragraphs
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # Extract index from paragraphs
    idx = None
    for p in paras:
        m = INDEX_PAT.search(p)
        if m:
            idx = m.group(0)
            break
    
    title = paras[0] if paras else rel
    
    # Extract tables
    for ti, table in enumerate(doc.tables):
        cells = []
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    cells.append({'r': ri, 'c': ci, 'v': text})
        
        if cells:
            # Try to find index in table cells too
            table_idx = idx
            for cell in cells:
                m = INDEX_PAT.match(cell['v'])
                if m and ('索引号' in str(cell['v']) or cell['r'] <= 1):
                    table_idx = m.group(0)
                    break
            
            entry = {
                'type': 'docx',
                'source': rel,
                'table_index': ti,
                'title': title,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'cells': cells,
            }
            if table_idx:
                entry['index_no'] = table_idx
            entries.append(entry)
    
    # Also add paragraphs as a separate entry if there's meaningful content
    meaningful_paras = [p for p in paras if len(p) > 10]
    if meaningful_paras and not entries:
        entries.append({
            'type': 'docx',
            'source': rel,
            'title': title,
            'paragraphs': meaningful_paras,
        })
    elif meaningful_paras:
        # Attach paragraphs to the first table entry
        entries[0]['paragraphs'] = meaningful_paras
    
    return entries

def main():
    files = get_files()
    print(f'找到 {len(files)} 个文件 ({sum(1 for _,e in files if e=="xlsx")} xlsx, {sum(1 for _,e in files if e=="docx")} docx)', flush=True)
    
    all_entries = OrderedDict()
    entry_counter = 0
    
    for fp, ext in files:
        rel = os.path.relpath(fp, BASE)
        print(f'\n处理: {rel}', flush=True)
        
        if ext == 'xlsx':
            entries = process_xlsx(fp, rel)
        elif ext == 'docx':
            entries = process_docx(fp, rel)
        else:
            continue
        
        for entry in entries:
            if 'error' in entry:
                print(f'  [错误] {entry.get("error","")}', flush=True)
                continue
            
            idx = entry.get('index_no', None)
            if idx:
                key = f'{idx}'
                if key in all_entries:
                    suffix = 2
                    while f'{idx}_{suffix}' in all_entries:
                        suffix += 1
                    key = f'{idx}_{suffix}'
                    print(f'  [!] 索引号重复: {idx}, 使用 key={key} 以避免覆盖', flush=True)
                all_entries[key] = entry
                entry_counter += 1
                print(f'  [{entry["type"]}] {key:12s} | {entry["source"]:45s} | {entry.get("title","")[:30]} | {len(entry.get("cells",[])):5d} cells', flush=True)
            else:
                # Entry without index - use a generated key
                key = f'_table_{entry_counter:04d}'
                all_entries[key] = entry
                entry_counter += 1
                print(f'  [{entry["type"]}] {key:12s} | {entry["source"]:45s} | {entry.get("title","")[:30]} | {len(entry.get("cells",[])):5d} cells [无索引号]', flush=True)
    
    # Write full JSON
    out = OrderedDict()
    for k in sorted(all_entries.keys(), key=lambda x: (all_entries[x].get('type',''), all_entries[x].get('source',''), x)):
        out[k] = all_entries[k]
    
    with open(OUTPUT, 'w', encoding='utf-8') as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    
    # Write compressed version (just index_no -> [title, cells])
    tmin = OrderedDict()
    for k, v in out.items():
        idx = v.get('index_no', k)
        if v.get('cells'):
            tmin[idx] = [v.get('title', ''), [[c['r'], c['c'], c['v']] for c in v['cells']]]
    
    with open(OUTPUT_MIN, 'w', encoding='utf-8') as f:
        json.dump(tmin, f, ensure_ascii=False)
    
    # Summary
    with_idx = sum(1 for v in out.values() if 'index_no' in v)
    without_idx = sum(1 for v in out.values() if 'index_no' not in v)
    
    print(f'\n{"="*60}', flush=True)
    print(f'✅ 完成!', flush=True)
    print(f'   总条目: {len(out)}', flush=True)
    print(f'   有索引号: {with_idx}', flush=True)
    print(f'   无索引号: {without_idx}', flush=True)
    print(f'   输出文件: {OUTPUT}', flush=True)
    print(f'   压缩版: {OUTPUT_MIN}', flush=True)

if __name__ == '__main__':
    main()
